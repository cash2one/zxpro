# coding:utf-8
from docx import Document

document = Document()
font = document.styles['Normal'].font
font.name = u"微软雅黑"
document.add_paragraph(u'编  号：15-1010')
document.add_paragraph(u'提案人：仲星、李新建')
document.add_paragraph(u'案  由：关于进一步优化漕河泾开发区交通组织运作的提案')
document.add_paragraph(u'主  办：区公安分局')
document.add_paragraph(u'会  办：虹梅街道')
document.add_paragraph(u'内  容：')
document.add_paragraph('A plain paragraph having some ')
document.add_paragraph('A plain paragraph having some ')
document.add_page_break()

document.save('demo1.docx')
