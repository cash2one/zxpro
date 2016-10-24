# -*-coding:utf-8 -*-
from django import forms
from django.core.exceptions import PermissionDenied
from django.db import router
from django.http import HttpResponse, HttpResponseRedirect
from django.template import loader
from django.template.response import TemplateResponse
from django.utils.datastructures import SortedDict
from django.utils.encoding import force_unicode
from django.utils.safestring import mark_safe
from django.utils.text import capfirst
from django.utils.translation import ugettext as _, ungettext
from xadmin.sites import site
from xadmin.util import model_format_dict, get_deleted_objects, model_ngettext
from xadmin.views import BaseAdminPlugin, ListAdminView
from xadmin.views.base import filter_hook, ModelAdminView

ACTION_CHECKBOX_NAME = '_selected_action'
checkbox = forms.CheckboxInput({'class': 'action-select'}, lambda value: False)


def action_checkbox(obj):
    return checkbox.render(ACTION_CHECKBOX_NAME, force_unicode(obj.pk))
action_checkbox.short_description = mark_safe(
    '<input type="checkbox" id="action-toggle" />')
action_checkbox.allow_tags = True
action_checkbox.allow_export = False
action_checkbox.is_column = False


class BaseActionView(ModelAdminView):
    action_name = None
    description = None
    icon = 'fa fa-tasks'

    model_perm = 'change'

    @classmethod
    def has_perm(cls, list_view):
        return list_view.get_model_perms()[cls.model_perm]

    def init_action(self, list_view):
        self.list_view = list_view
        self.admin_site = list_view.admin_site

    @filter_hook
    def do_action(self, queryset):
        pass


from xhzx.models import Jianfenchi,Jiafenchi,MnTTa,AutoSubmit,MnTWy,MnTMeetingType,MnTMeeting,MnTMeetingDetail,MnTZwh
class DeleteSelectedAction(BaseActionView):

    action_name = "delete_selected"
    description = _(u'Delete selected %(verbose_name_plural)s')

    delete_confirmation_template = None
    delete_selected_confirmation_template = None

    delete_models_batch = True

    model_perm = 'delete'
    icon = 'fa fa-times'

    @filter_hook
    def delete_models(self, queryset):
        n = queryset.count()
        if n:
            if self.delete_models_batch:
                for obj in queryset:
                    # print (type(obj) == MnTTa)
                    # if obj.mno and "," in obj.mno:
                    #     type = 2
                    # else:
                    #     type = 1
                    # Jianfenchi.objects.create(djf_issueno=obj.issueno,
                    #                           djf_issueid=obj.pk,
                    #                           djf_wyno=obj.mno,
                    #                           djf_issuebt=obj.biaoti,
                    #                           djf_wyname=obj.mname,
                    #                           djf_type=type,
                    #                           djf_time=obj.submit_time)
                    obj.delete()
            else:
                for obj in queryset:
                    obj.delete()
            self.message_user(_("Successfully deleted %(count)d %(items)s.") % {
                "count": n, "items": model_ngettext(self.opts, n)
            }, 'success')

    @filter_hook
    def do_action(self, queryset):
        # Check that the user has delete permission for the actual model
        if not self.has_delete_permission():
            raise PermissionDenied

        using = router.db_for_write(self.model)

        # Populate deletable_objects, a data structure of all related objects that
        # will also be deleted.
        deletable_objects, perms_needed, protected = get_deleted_objects(
            queryset, self.opts, self.user, self.admin_site, using)

        # The user has already confirmed the deletion.
        # Do the deletion and return a None to display the change list view again.
        if self.request.POST.get('post'):
            if perms_needed:
                raise PermissionDenied
            self.delete_models(queryset)
            # Return None to display the change list page again.
            return None

        if len(queryset) == 1:
            objects_name = force_unicode(self.opts.verbose_name)
        else:
            objects_name = force_unicode(self.opts.verbose_name_plural)

        if perms_needed or protected:
            title = _("Cannot delete %(name)s") % {"name": objects_name}
        else:
            title = _("Are you sure?")

        context = self.get_context()
        context.update({
            "title": title,
            "objects_name": objects_name,
            "deletable_objects": [deletable_objects],
            'queryset': queryset,
            "perms_lacking": perms_needed,
            "protected": protected,
            "opts": self.opts,
            "app_label": self.app_label,
            'action_checkbox_name': ACTION_CHECKBOX_NAME,
        })

        # Display the confirmation page
        return TemplateResponse(self.request, self.delete_selected_confirmation_template or
                                self.get_template_list('views/model_delete_selected_confirm.html'), context, current_app=self.admin_site.name)


class ActionPlugin(BaseAdminPlugin):

    # Actions
    actions = []
    actions_selection_counter = True
    global_actions = [DeleteSelectedAction]

    def init_request(self, *args, **kwargs):
        self.actions = self.get_actions()
        return bool(self.actions)

    def get_list_display(self, list_display):
        if self.actions:
            list_display.insert(0, 'action_checkbox')
            self.admin_view.action_checkbox = action_checkbox
        return list_display

    def get_list_display_links(self, list_display_links):
        if self.actions:
            if len(list_display_links) == 1 and list_display_links[0] == 'action_checkbox':
                return list(self.admin_view.list_display[1:2])
        return list_display_links

    def get_context(self, context):
        if self.actions and self.admin_view.result_count:
            av = self.admin_view
            selection_note_all = ungettext('%(total_count)s selected',
                                           'All %(total_count)s selected', av.result_count)

            new_context = {
                'selection_note': _('0 of %(cnt)s selected') % {'cnt': len(av.result_list)},
                'selection_note_all': selection_note_all % {'total_count': av.result_count},
                'action_choices': self.get_action_choices(),
                'actions_selection_counter': self.actions_selection_counter,
            }
            context.update(new_context)
        return context

    def post_response(self, response, *args, **kwargs):
        request = self.admin_view.request
        av = self.admin_view

        # Actions with no confirmation
        if self.actions and 'action' in request.POST:
            action = request.POST['action']

            if action not in self.actions:
                msg = _("Items must be selected in order to perform "
                        "actions on them. No items have been changed.")
                av.message_user(msg)
            else:
                ac, name, description, icon = self.actions[action]
                select_across = request.POST.get('select_across', False) == '1'
                selected = request.POST.getlist(ACTION_CHECKBOX_NAME)

                if not selected and not select_across:
                    # Reminder that something needs to be selected or nothing will happen
                    msg = _("Items must be selected in order to perform "
                            "actions on them. No items have been changed.")
                    av.message_user(msg)
                else:
                    queryset = av.list_queryset._clone()
                    if not select_across:
                        # Perform the action only on the selected objects
                        queryset = av.list_queryset.filter(pk__in=selected)
                    response = self.response_action(ac, queryset)
                    # Actions may return an HttpResponse, which will be used as the
                    # response from the POST. If not, we'll be a good little HTTP
                    # citizen and redirect back to the changelist page.
                    if isinstance(response, HttpResponse):
                        return response
                    else:
                        return HttpResponseRedirect(request.get_full_path())
        return response

    def response_action(self, ac, queryset):
        if isinstance(ac, type) and issubclass(ac, BaseActionView):
            action_view = self.get_model_view(ac, self.admin_view.model)
            action_view.init_action(self.admin_view)
            return action_view.do_action(queryset)
        else:
            return ac(self.admin_view, self.request, queryset)

    def get_actions(self):
        if self.actions is None:
            return SortedDict()

        actions = [self.get_action(action) for action in self.global_actions]

        for klass in self.admin_view.__class__.mro()[::-1]:
            class_actions = getattr(klass, 'actions', [])
            if not class_actions:
                continue
            actions.extend(
                [self.get_action(action) for action in class_actions])

        # get_action might have returned None, so filter any of those out.
        actions = filter(None, actions)

        # Convert the actions into a SortedDict keyed by name.
        actions = SortedDict([
            (name, (ac, name, desc, icon))
            for ac, name, desc, icon in actions
        ])

        return actions

    def get_action_choices(self):
        """
        Return a list of choices for use in a form object.  Each choice is a
        tuple (name, description).
        """
        choices = []
        for ac, name, description, icon in self.actions.itervalues():
            choice = (name, description % model_format_dict(self.opts), icon)
            choices.append(choice)
        return choices

    def get_action(self, action):
        if isinstance(action, type) and issubclass(action, BaseActionView):
            if not action.has_perm(self.admin_view):
                return None
            return action, getattr(action, 'action_name'), getattr(action, 'description'), getattr(action, 'icon')

        elif callable(action):
            func = action
            action = action.__name__

        elif hasattr(self.admin_view.__class__, action):
            func = getattr(self.admin_view.__class__, action)

        else:
            return None

        if hasattr(func, 'short_description'):
            description = func.short_description
        else:
            description = capfirst(action.replace('_', ' '))

        return func, action, description, getattr(func, 'icon', 'tasks')

    # View Methods
    def result_header(self, item, field_name, row):
        if item.attr and field_name == 'action_checkbox':
            item.classes.append("action-checkbox-column")
        return item

    def result_item(self, item, obj, field_name, row):
        if item.field is None and field_name == u'action_checkbox':
            item.classes.append("action-checkbox")
        return item

    # Media
    def get_media(self, media):
        if self.actions and self.admin_view.result_count:
            media = media + self.vendor('xadmin.plugin.actions.js', 'xadmin.plugins.css')
        return media

    # Block Views
    def block_results_bottom(self, context, nodes):
        if self.actions and self.admin_view.result_count:
            nodes.append(loader.render_to_string('xadmin/blocks/model_list.results_bottom.actions.html', context_instance=context))


from xhzx.models import MnTBljg_sqmy, Jiafenchi


class act_fabu(BaseActionView):


    action_name = "public"
    description = _(u'发布 %(verbose_name_plural)s')

    model_perm = 'change'


    def do_action(self, queryset):
        import datetime

        if queryset[0] and queryset[0].class_field == "提案":
            for obj in queryset:
                obj.publish = True
                obj.admin_status_id = 1
                obj.publish_time = datetime.datetime.now()
                obj.save()
                Jiafenchi.objects.create(djf_issueno=obj.issueno,
                                         djf_tos=obj.class_field,
                                         djf_issueid=obj.pk,
                                         djf_issuebt=obj.biaoti,
                                         djf_wyno=obj.mno,
                                         djf_wyname=obj.mname,
                                         djf_state=2,
                                         djf_type=2 if (obj.mno and ',' in obj.mno) else 1,
                                         djf_time=obj.submit_time)

        else:
            for obj in queryset:
                obj.publish = True

                zzbl = MnTBljg_sqmy.objects.get(value=1)
                obj.admin_status.add(zzbl)
                obj.publish_time = datetime.datetime.now()
                obj.save()
        self.message_user("操作成功!","success")



class cleardafu(BaseActionView):


    action_name = "cleardafu"
    description = _(u'清除答复')
    model_perm = 'change'


    def do_action(self, queryset):
        buchenggong = ""
        for obj in queryset:
            try:
                obj.banlidafu = ""
                obj.yidafu = False
                obj.banlijieguo = ""
                obj.answer_time = None
                obj.save()
            except Exception as e:
                buchenggong += str(e)

        self.message_user("设置完成:%s" % buchenggong,"success")

from datetime import datetime
class autosubmit(BaseActionView):
    action_name = "autosubmit"
    description = _(u'自动上传 %(verbose_name_plural)s')

    model_perm = 'change'

    def do_action(self, queryset):
        fb = {'data': True, 'msg': ""}
        count = 0
        ecount = 0
        xxxxx = []
        ccount = 0
        scount = 0
        for  ta in queryset:
            count += 1
            fb['data'] = True
            fb['msg'] = ta.issueno if ta.issueno else ""

            test = AutoSubmit.objects.filter(pro_tano=ta.issueno)
            if test.count() > 0:
                ccount += 1
                continue
            if not ta.issueno:
                fb['msg'] += "提案编号不能为空 "
                fb['data'] = False
            if not ta.biaoti:
                fb['msg'] += "提案标题不能为空 "
                fb['data'] = False


            if (not ta.yjhjy) and (not ta.qkfy):
                fb['msg'] += "情况反映 或 意见和建议 不能全为空 "
                fb['data'] = False

            if not ta.reconfirmnames:
                fb['msg'] += "分类不能为空 "
                fb['data'] = False
            if ta.zhuban.all().count() and not ta.heban.all().count():
                pass
            elif ta.heban.all().count() and not ta.zhuban.all().count():
                pass
            else:
                fb['msg'] += "主办或合办只能选择一个 "
                fb['data'] = False
            if fb['data']:
                try:
                    if " " in ta.writer:
                        writer = (ta.writer.split())[0]
                    else:
                        writer = ta.writer

                    if "-" in writer:
                        writerwy = (MnTWy.objects.filter(wyno=writer))[0]
                    else:
                        writerwy = (MnTWy.objects.filter(name=writer))[0]


                    fl = ta.reconfirmnames
                    if fl == "A":
                        leixing = "财政经济类"
                    elif fl == "B":
                        leixing = "城市建设和管理类"
                    elif fl == "C":
                        leixing = "教科文卫体类"
                    elif fl == "D":
                        leixing = "综合类"
                    else:
                        leixing = ta.reconfirmnames

                    try:
                        wyzwh = writerwy.zwh.text
                    except:
                        wyzwh = ""

                    try:
                        wyjiebie = writerwy.jiebie.text
                    except:
                        wyjiebie = ""
                    try:
                        wyparty = writerwy.party.text
                    except:
                        wyparty = ""

                    a = AutoSubmit.objects.create(
                            pro_tano= ta.issueno,
                            pro_title= ta.biaoti,
                            pro_qkfy= ta.qkfy,
                            pro_yjhjy= ta.yjhjy,
                        pro_cbdw=",".join(map(lambda x: x.text, ta.zhuban.all())),
                            pro_type= "个人联名" if ta.grper == "个人" else "党派团体",
                            pro_acceptdate= str(datetime.strftime(ta.submit_time, "%Y-%m-%d")),
                        pro_hebandw=",".join(map(lambda x: x.text, ta.heban.all())),
                        pro_huibandw=",".join(map(lambda x: x.text, ta.huiban.all())),
                            pro_leader= ta.fenguan if ta.fenguan else "",
                        pro_bo_type=leixing,
                            pro_partyindex= "区政协%s会议" % ta.hyjc,
                            zz_name= ta.mname,
                            zz_tell1= ("%s" % writerwy.workphone) if writerwy.workphone else "",
                            zz_tell2= writerwy.cellphone if writerwy.cellphone else "",
                            zz_gzdw= writerwy.workpos if writerwy.workpos else "",
                            zz_zwh= wyzwh,
                            zz_jb= wyjiebie,
                            zz_dp= wyparty
                    )
                    if a:
                        fb['data'] = True
                        fb['msg'] = ''
                        scount += 1
                    else:
                        fb['data'] = False
                        fb['msg'] += '添加失败'
                        xxxxx.append(str(ta.pk))
                        ecount += 1
                except Exception as e:
                    fb['data'] = False
                    fb['msg'] += '添加失败:%s' % str(e)
                    xxxxx.append(str(ta.pk))
                    ecount += 1
                print "%s:%s" % (ta.issueno ,fb['msg'])

            else:
                xxxxx.append(str(ta.pk))
                ecount += 1
                continue
        self.message_user("共计:%s,成功:%s,失败:%s[%s],跳过:%s\n" % (count,scount,ecount,",".join(xxxxx),ccount),"success")


class act_ta_download(BaseActionView):
    action_name = "download_hz"
    description = _(u'下载 %(verbose_name_plural)s 汇总')
    model_perm = 'change'

    def do_action(self, queryset):

        import os, csv
        file_name = "tian.csv"
        try:
            if os.path.exists(file_name):
                os.remove(file_name)
        except:
            pass
        with open(file_name, 'wb') as csvfile:
            spamwriter = csv.writer(csvfile, dialect='excel')
            spamwriter.writerow([""])
            spamwriter.writerow(["编  号", "提案人", "案  由", "类别", "主  办", "会  办", "办理情况"])
            for ta in queryset:
                spamwriter.writerow([ta.issueno, ta.mname, ta.biaoti, ta.reconfirmnames, ",".join(map(lambda x: x.text, ta.zhuban.all())),",".join(map(lambda x: x.text, ta.huiban.all())), ta.admin_status.text if ta.admin_status else ""])

        return HttpResponseRedirect('/get/tian/')

class act_sq_download(BaseActionView):
    action_name = "download_sqhz"
    description = _(u'下载 %(verbose_name_plural)s excel')
    model_perm = 'change'

    def do_action(self, queryset):

        import os, csv
        file_name = "sqmy.csv"
        try:
            if os.path.exists(file_name):
                os.remove(file_name)
        except:
            pass
        with open(file_name, 'wb') as csvfile:
            spamwriter = csv.writer(csvfile, dialect='excel')
            spamwriter.writerow([""])
            spamwriter.writerow(["编  号", "作  者", "标  题", "办理情况", "", "", ""])
            for ta in queryset:
                spamwriter.writerow([ta.issueno, ta.mname, ta.biaoti, ",".join(map(lambda x: x.text, ta.admin_status.all())), "","",""])

        return HttpResponseRedirect('/get/sqmy/')


class act_tadoc_download(BaseActionView):
    action_name = "download_td"
    description = _(u'下载 %(verbose_name_plural)s doc')
    model_perm = 'change'

    def do_action(self, queryset):
        import os, shutil
        dirname = "tar"
        try:
            if os.path.exists(dirname):
                shutil.rmtree("%s/" % dirname)
                os.mkdir(dirname)
            else:
                os.mkdir(dirname)
        except:
            pass
        try:
            if os.path.exists('tar.zip'):
                os.remove('tar.zip')
        except:
            pass
        from docx import Document
        for ta in queryset:
            try:
                document = Document()
                document.add_paragraph(u'编  号:%s' % ta.issueno)
                document.add_paragraph(u'提案人:%s' % ta.mname)
                document.add_paragraph(u'案  由:%s' % ta.biaoti)
                document.add_paragraph(u'主  办:')
                document.add_paragraph(u'会  办:')
                document.add_paragraph(u'内  容:')
                document.add_paragraph(u"%s" % ta.qkfy)
                document.add_paragraph(u"%s" % ta.yjhjy)
                document.add_page_break()
                document.save("tar/%s.docx" % (ta.issueno if ta.issueno else ta.pk))
            except:
                continue

        import zipfile
        try:
            import zlib
            compression = zipfile.ZIP_DEFLATED
        except:
            compression = zipfile.ZIP_STORED
        path = 'tar/'  # 要进行压缩的文档目录
        start = path.rfind(os.sep) + 1
        filename = 'tar.zip'  # 压缩后的文件名
        z = zipfile.ZipFile(filename, mode="w", compression=compression)
        try:
            for dirpath, dirs, files in os.walk(path):
                for file in files:
                    if file == filename:
                        continue
                    print(file)
                    z_path = os.path.join(dirpath, file)
                    z.write(z_path, z_path[start:])
            z.close()
        except:
            if z:
                z.close()
        try:
            shutil.copy('tar.zip', '/static/tar.zip')
        except Exception as e:
            print e

        return HttpResponseRedirect('/get/tazip/')



class act_meeting_download(BaseActionView):
    action_name = "download_td"
    description = _(u'下载 %(verbose_name_plural)s 汇总')
    model_perm = 'change'
    wys = MnTWy.objects.all()
    def GetNameByWyno(self,wyno):
        wyname = '查无此人'
        for w in self.wys:
            if w.wyno == wyno:
                wyname = w.name
        return wyname

    def GetWyInZwh(self,zwh,wynos):
        wynames = []
        for owy in wynos:
            try:
                w = MnTWy.objects.get(wyno=owy)
                if w.zwh == zwh:
                    wynames.append(w.name)
                elif w.zwh2 == zwh:
                    wynames.append(w.name)
                else:
                    pass
            except:
                pass


        return ','.join(wynames)



    def do_action(self, queryset):
        import os, shutil
        dirname = "tar"
        try:
            if os.path.exists(dirname):
                shutil.rmtree("%s/" % dirname)
                os.mkdir(dirname)
            else:
                os.mkdir(dirname)
        except:
            pass
        try:
            if os.path.exists('tar.zip'):
                os.remove('tar.zip')
        except:
            pass

        import os, csv
        file_name = "tian.csv"
        try:
            if os.path.exists(file_name):
                os.remove(file_name)
        except:
            pass
        with open(file_name, 'wb') as csvfile:
            spamwriter = csv.writer(csvfile, dialect='excel')
            spamwriter.writerow([""])
            spamwriter.writerow(["会议时间","会议名称","会议地点","应到","实到","未到"])

            for omeeting in queryset:
                yingdao = ",".join(map(lambda x: self.GetNameByWyno(x.wyno), MnTMeetingDetail.objects.filter(mid=omeeting.pk)))
                shidao = ",".join(map(lambda x: self.GetNameByWyno(x.wyno), MnTMeetingDetail.objects.filter(mid=omeeting.pk,ycj=1)))
                weidao = ",".join(map(lambda x: self.GetNameByWyno(x.wyno), MnTMeetingDetail.objects.filter(mid=omeeting.pk,yqx=1)))
                spamwriter.writerow([omeeting.time,omeeting.no,omeeting.place,yingdao,shidao,weidao])

        for ozwh in MnTZwh.objects.all():

            file_name = "%s.csv" % ozwh.text
            try:
                if os.path.exists(file_name):
                    os.remove(file_name)
            except:
                pass
            with open(file_name, 'wb') as csvfile:
                spamwriter = csv.writer(csvfile, dialect='excel')
                spamwriter.writerow([""])
                spamwriter.writerow(["会议时间","会议名称","会议地点","应到","实到","未到"])

                for omeeting in queryset:
                    yingdao = self.GetWyInZwh(ozwh,MnTMeetingDetail.objects.filter(mid=omeeting.pk))
                    shidao = self.GetWyInZwh(ozwh,MnTMeetingDetail.objects.filter(mid=omeeting.pk,ycj=1))
                    weidao = self.GetWyInZwh(ozwh,MnTMeetingDetail.objects.filter(mid=omeeting.pk,yqx=1))
                    spamwriter.writerow([omeeting.time,omeeting.no,omeeting.place,yingdao,shidao,weidao])


        import zipfile
        try:
            import zlib
            compression = zipfile.ZIP_DEFLATED
        except:
            compression = zipfile.ZIP_STORED
        path = 'tar/'  # 要进行压缩的文档目录
        start = path.rfind(os.sep) + 1
        filename = 'tar.zip'  # 压缩后的文件名
        z = zipfile.ZipFile(filename, mode="w", compression=compression)
        try:
            for dirpath, dirs, files in os.walk(path):
                for file in files:
                    if file == filename:
                        continue
                    print(file)
                    z_path = os.path.join(dirpath, file)
                    z.write(z_path, z_path[start:])
            z.close()
        except:
            if z:
                z.close()
        try:
            shutil.copy('tar.zip', '/static/tar.zip')
        except Exception as e:
            print e

        return HttpResponseRedirect('/get/tazip/')


class act_sqdoc_download(BaseActionView):
    action_name = "download_td"
    description = _(u'下载 %(verbose_name_plural)s doc')
    model_perm = 'change'

    def do_action(self, queryset):
        import os, shutil
        dirname = "tar"
        try:
            if os.path.exists(dirname):
                shutil.rmtree("%s/" % dirname)
                os.mkdir(dirname)
            else:
                os.mkdir(dirname)
        except:
            pass
        try:
            if os.path.exists('tar.zip'):
                os.remove('tar.zip')
        except:
            pass
        from docx import Document

        for ta in queryset:
            try:
                document = Document()
                document.add_paragraph(u'编  号:%s' % ta.issueno)
                document.add_paragraph(u'反映者:%s' % ta.mname)
                document.add_paragraph(u'题  目:%s' % ta.biaoti)
                document.add_paragraph(u'办理结果:')
                document.add_paragraph(u"%s" % ta.qkfy)
                document.add_paragraph(u"%s" % ta.yjhjy)
                document.add_page_break()
                document.save("tar/%s.docx" % (ta.issueno if ta.issueno else ta.pk))
            except:
                continue

        import zipfile
        try:
            import zlib
            compression = zipfile.ZIP_DEFLATED
        except:
            compression = zipfile.ZIP_STORED
        path = 'tar/'  # 要进行压缩的文档目录
        start = path.rfind(os.sep) + 1
        filename = 'tar.zip'  # 压缩后的文件名
        z = zipfile.ZipFile(filename, mode="w", compression=compression)
        try:
            for dirpath, dirs, files in os.walk(path):
                for file in files:
                    if file == filename:
                        continue
                    print(file)
                    z_path = os.path.join(dirpath, file)
                    z.write(z_path, z_path[start:])
            z.close()
        except:
            if z:
                z.close()
        try:
            # shutil.copy('tar.zip', 'C:/xhzxproadmin/xhzx/static/tar.zip')
            shutil.copy('tar.zip', '/static/tar.zip')
        except:
            pass

        return HttpResponseRedirect('/get/tazip/')


class act_sqctdoc_download(BaseActionView):
    action_name = "download_td"
    description = _(u'下载 %(verbose_name_plural)s 打印内容(doc文档)')
    model_perm = 'change'

    def do_action(self, queryset):
        import os, shutil
        dirname = "tar"
        try:
            if os.path.exists(dirname):
                shutil.rmtree("%s/" % dirname)
                os.mkdir(dirname)
            else:
                os.mkdir(dirname)
        except:
            pass
        try:
            if os.path.exists('tar.zip'):
                os.remove('tar.zip')
        except:
            pass
        from docx import Document
        for ta in queryset:
            document = Document()
            document.add_paragraph(u'中国人民政治协商会议上海市徐汇区委员会')
            document.add_paragraph(u'社情民意第%s号' % ta.issueno)
            document.add_paragraph(u'%s' % ta.submit_time)

            from xhzx.models import MnTWy

            if ta.sqmy_zbname:
                name = ta.writer
                cellphone = ta.sqmy_contact
                zbwork = ta.sqmy_zbwork
            else:

                if ('-' in ta.writer) or ('0000' in ta.writer):
                    wr = MnTWy.objects.get(wyno=ta.writer)
                    name = wr.name
                    cellphone = wr.cellphone
                    zbwork = "%s %s" % (wr.workplace, wr.workpos)

                else:
                    try:
                        wr = MnTWy.objects.get(name=ta.writer)
                        name = wr.name
                        cellphone = wr.cellphone
                        zbwork = "%s %s" % (wr.workplace, wr.workpos)

                    except:
                        name = ta.writer
                        cellphone = ""
                        zbwork = ""
            document.add_paragraph(u'反映单位或反映人:%s' % name)
            document.add_paragraph(u'联系电话:%s' % cellphone)
            document.add_paragraph(u'工作单位和职务:%s' % zbwork)
            document.add_paragraph(u'标题:%s' % ta.biaoti)

            document.add_paragraph(u'情况反映:')

            document.add_paragraph(ta.qkfy)
            document.add_paragraph(u'建议:')
            document.add_paragraph(ta.yjhjy)
            document.add_page_break()
            document.save("tar/%s.docx" % (ta.issueno if ta.issueno else ta.pk))

        import zipfile
        try:
            import zlib
            compression = zipfile.ZIP_DEFLATED
        except:
            compression = zipfile.ZIP_STORED
        path = 'tar/'  # 要进行压缩的文档目录
        start = path.rfind(os.sep) + 1
        filename = 'tar.zip'  # 压缩后的文件名
        z = zipfile.ZipFile(filename, mode="w", compression=compression)
        try:
            for dirpath, dirs, files in os.walk(path):
                for file in files:
                    if file == filename:
                        continue
                    print(file)
                    z_path = os.path.join(dirpath, file)
                    z.write(z_path, z_path[start:])
            z.close()
        except:
            if z:
                z.close()
        try:
            shutil.copy('tar.zip', '/static/tar.zip')
        except Exception as  e:
            print e

        return HttpResponseRedirect('/get/tazip/')
site.register_plugin(ActionPlugin, ListAdminView)


